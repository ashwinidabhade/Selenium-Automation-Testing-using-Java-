from docx import Document
import os

def generate_document(data):
    os.makedirs("output", exist_ok=True)
    doc = Document()
    doc.add_heading('Performance Qualification Report', 0)
    doc.add_paragraph(f"Customer Name: {data['customer_name']}")
    doc.add_paragraph(f"Address: {data['address']}")
    doc.add_paragraph(f"Email: {data['email']}")
    doc.add_paragraph(f"Objective: {data['objective']}")
    doc.add_paragraph(f"Calibrated Value: {data['calibrated_value']}")
    doc.add_paragraph(f"Desired Value: {data['desired_value']}")

    filename = f"output/PQ_Report_{data['customer_name'].replace(' ', '_')}.docx"
    doc.save(filename)
    print("✅ PQ document generated:", filename)
